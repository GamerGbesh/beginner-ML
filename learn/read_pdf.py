import fitz

from docx import Document
import pytesseract
from PIL import Image


# For reading text heavy pdfs and writing it to a word file
document = Document()

title = "Mensah Philemon's CV"
doc = fitz.open(title + ".pdf")
document.add_heading(title, level=1)
for page in doc:
    text = page.get_text()
    document.add_paragraph(text)

document.save("my_document.docx")


# For reading scanned pdfs containing texts
document2 = Document()
document2.add_heading("Pasco", 0)
scanned_doc = fitz.open("Database pasco1.pdf")
for page in scanned_doc:
    image_list = page.get_images(full=True)
    for image in image_list:
        xref = image[0]
        image_pil = fitz.Pixmap(scanned_doc, xref).pil_image()
        text_from_image = pytesseract.image_to_string(image_pil)
        document2.add_paragraph(text_from_image)

document2.save("Passco.docx")